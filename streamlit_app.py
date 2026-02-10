#!/usr/bin/env python3
"""
Two-Pass Reformulation Detection System
Pass 1: Small model (gpt-oss:20b) - Liberal detection
Pass 2: Large model (deepseek-r1 or llama3:70b) - Strict validation
"""

import streamlit as st
import ollama
from ollama import Client
import os
import json
import re
from pathlib import Path
from typing import List, Dict, Tuple
import time

try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    from io import BytesIO
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Initialize Ollama Cloud client from secrets
client = None
if "ollama_key" in st.secrets:
    try:
        client = Client(
            host="https://ollama.com",
            headers={'Authorization': 'Bearer ' + st.secrets["ollama_key"]}
        )
    except Exception as e:
        print(f"Error initializing client from secrets: {e}")

# Page config
st.set_page_config(
    page_title="Reformulation Detector (Two-Pass)",
    page_icon="🔍",
    layout="wide"
)

# Get the directory where this script is located
SCRIPT_DIR = Path(__file__).parent

def load_prompt(filename: str, secret_key: str, default: str = "") -> str:
    """Load a prompt from streamlit secrets (if available) or file."""
    # Try secret first
    if secret_key in st.secrets:
        return st.secrets[secret_key]

    # Fallback to file
    prompt_path = SCRIPT_DIR / filename
    try:
        return prompt_path.read_text(encoding='utf-8')
    except int:
        pass # Ignore file errors if not found, will return default
    except Exception as e:
        st.warning(f"⚠️ Error loading '{filename}': {e}. Using default prompt.")
    
    return default

# Load prompts
DETECTION_PROMPT = load_prompt("detection_prompt.txt", "detection_prompt", "Detect reformulation pairs and trios in child-adult conversations.")
VALIDATION_PROMPT = load_prompt("validation_prompt.txt", "validation_prompt", "Validate detected reformulations and keep only genuine ones.")


def parse_transcript(file_content: str, file_type: str) -> List[Tuple[int, str]]:
    """Parse transcript content into (line_number, utterance) tuples."""
    lines = []
    line_counter = 1
    
    # Parse text content
    for line in file_content.strip().split('\n'):
        line = line.strip()
        if not line:
            continue
        
        # Try format 1: number. SPEAKER: utterance
        match = re.match(r'^(\d+)\.\s+([A-Z]+):\s+(.+)$', line)
        if match:
            line_num = int(match.group(1))
            speaker = match.group(2)
            utterance = match.group(3)
            lines.append((line_num, f"{speaker}: {utterance}"))
            continue
        
        # Try format 2: SPEAKER: utterance (no number)
        match = re.match(r'^([A-Z]+):\s+(.+)$', line)
        if match:
            speaker = match.group(1)
            utterance = match.group(2)
            lines.append((line_counter, f"{speaker}: {utterance}"))
            line_counter += 1
    
    return lines


def create_chunks(lines: List[Tuple[int, str]], chunk_size: int = 20, overlap: int = 3) -> List[List[Tuple[int, str]]]:
    """Split transcript into chunks with overlap."""
    if overlap >= chunk_size:
        overlap = chunk_size - 1
    
    chunks = []
    i = 0
    
    while i < len(lines):
        chunk = lines[i:i + chunk_size]
        chunks.append(chunk)
        i += (chunk_size - overlap)
    
    return chunks


def create_analysis_prompt(chunk: List[Tuple[int, str]], pass_type: str = "detection") -> str:
    """Create prompt for analyzing a chunk."""
    formatted = [f"{num}. {text}" for num, text in chunk]
    chunk_text = "\n".join(formatted)
    
    if pass_type == "detection":
        return f"""Analyze this conversation chunk for potential reformulation pairs and trios.

CONVERSATION LINES:
{chunk_text}

TASK: Cast a wide net - include ALL potential reformulations, even uncertain ones.
Use confidence tags: HIGH, MEDIUM, or LOW.

Return ONLY valid JSON with your findings (no markdown, no extra text)."""
    else:  # validation
        return chunk_text


def create_validation_prompt(detected_items: Dict, chunk_context: List[Tuple[int, str]] = None) -> str:
    """Create prompt for validating detected reformulations with optional context."""
    context_text = ""
    if chunk_context:
        context_text = "\n\nCONVERSATION CONTEXT:\n"
        context_text += "\n".join([f"{num}. {text}" for num, text in chunk_context])
    
    return f"""Review these detected reformulations and keep only genuine ones.

DETECTED REFORMULATIONS:
{json.dumps(detected_items, indent=2)}
{context_text}

TASK: For each pair/trio, decide KEEP or REJECT.
Add "validation_status" and "validation_reason"/"rejection_reason" fields.

Return JSON with validated pairs, trios, and rejected items."""


def query_ollama_stream(prompt: str, model: str, system_prompt: str, stream_placeholder=None):
    """Query Ollama with streaming support and return thinking + content."""
    try:
        full_response = ""
        thinking_content = ""
        has_thinking = False
        
        stream = client.chat(
            model=model,
            messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': prompt}
            ],
            options={
                'temperature': 0.1,
                'num_predict': 4096,
            },
            stream=True
        )
        
        # Stream response in real-time
        for chunk_response in stream:
            if "message" in chunk_response:
                # Capture thinking if available
                if "thinking" in chunk_response["message"] and chunk_response["message"]["thinking"]:
                    thinking_content += chunk_response["message"]["thinking"]
                    has_thinking = True
                
                # Capture response content
                if "content" in chunk_response["message"] and chunk_response["message"]["content"]:
                    full_response += chunk_response["message"]["content"]
                
                # Update display if placeholder provided
                if stream_placeholder:
                    with stream_placeholder.container():
                        if has_thinking:
                            with st.expander("💭 Model's Thinking", expanded=True):
                                st.markdown(thinking_content)
                        with st.expander("📝 Model's Response", expanded=True):
                            st.code(full_response, language="json")
        
        return {
            'content': full_response,
            'thinking': thinking_content if has_thinking else None,
            'success': True
        }
    except Exception as e:
        return {
            'content': None,
            'thinking': None,
            'success': False,
            'error': str(e) + "\n\n(Note: My quota may be exceeded. Please bring your own key at https://ollama.com/settings/keys)"
        }


def extract_json(response: str) -> Dict:
    """Extract and parse JSON from model response."""
    response = re.sub(r'```json\s*', '', response)
    response = re.sub(r'```\s*', '', response)
    response = response.strip()
    
    json_match = re.search(r'\{.*\}', response, re.DOTALL)
    if json_match:
        try:
            return json.loads(json_match.group(0))
        except json.JSONDecodeError:
            return {"pairs": [], "trios": []}
    
    return {"pairs": [], "trios": []}


def create_highlighted_docx(lines: List[Tuple[int, str]], pairs: List[Dict], trios: List[Dict]) -> Document:
    """Create a DOCX document with pairs highlighted in yellow and trios in green."""
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    doc.add_heading('Reformulation Analysis Report (Validated)', 0)
    
    # Add summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f'Total Lines: {len(lines)}')
    doc.add_paragraph(f'Validated Pairs: {len(pairs)}')
    doc.add_paragraph(f'Validated Trios: {len(trios)}')
    
    # Show confidence breakdown
    high_count = sum(1 for p in pairs + trios if p.get('confidence') == 'HIGH')
    medium_count = sum(1 for p in pairs + trios if p.get('confidence') == 'MEDIUM')
    low_count = sum(1 for p in pairs + trios if p.get('confidence') == 'LOW')
    
    doc.add_paragraph(f'Confidence: {high_count} HIGH, {medium_count} MEDIUM, {low_count} LOW')
    doc.add_paragraph('')
    
    # Create lookup dictionaries for highlighting
    line_highlights = {}
    
    # Process pairs (yellow highlight)
    for i, pair in enumerate(pairs):
        child_line = pair['child_line']
        adult_line = pair['adult_line']
        conf = pair.get('confidence', 'UNKNOWN')
        line_highlights[child_line] = ('yellow', f'Pair {i+1} [{conf}] - Child Error', i)
        line_highlights[adult_line] = ('yellow', f'Pair {i+1} [{conf}] - Adult Reformulation', i)
    
    # Process trios (green highlight - overrides pairs if same line)
    for i, trio in enumerate(trios):
        child_line = trio['child_line']
        adult_line = trio['adult_line']
        repair_line = trio['repair_line']
        conf = trio.get('confidence', 'UNKNOWN')
        line_highlights[child_line] = ('green', f'Trio {i+1} [{conf}] - Child Error', i)
        line_highlights[adult_line] = ('green', f'Trio {i+1} [{conf}] - Adult Reformulation', i)
        line_highlights[repair_line] = ('green', f'Trio {i+1} [{conf}] - Child Self-Repair', i)
    
    # Add legend
    doc.add_heading('Legend', level=1)
    legend_para = doc.add_paragraph()
    yellow_run = legend_para.add_run('■ Yellow Highlight')
    yellow_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    legend_para.add_run(' = Reformulation Pair (Child Error → Adult Correction)')
    
    legend_para2 = doc.add_paragraph()
    green_run = legend_para2.add_run('■ Green Highlight')
    green_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    legend_para2.add_run(' = Reformulation Trio (Child Error → Adult Correction → Child Self-Repair)')
    
    doc.add_paragraph('Confidence: [HIGH], [MEDIUM], or [LOW]')
    doc.add_paragraph('')
    
    # Add transcript with highlights
    doc.add_heading('Annotated Transcript', level=1)
    
    for line_num, text in lines:
        para = doc.add_paragraph()
        
        # Add line number and text
        full_text = f"{line_num}. {text}"
        run = para.add_run(full_text)
        
        # Apply highlighting if this line is part of a pair or trio
        if line_num in line_highlights:
            color, role, index = line_highlights[line_num]
            if color == 'yellow':
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif color == 'green':
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            # Add annotation
            annotation = para.add_run(f'  [{role}]')
            annotation.font.italic = True
            annotation.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph('')
    doc.add_page_break()
    
    # Add detailed pair analysis
    if pairs:
        doc.add_heading('Detailed Pair Analysis', level=1)
        for i, pair in enumerate(pairs, 1):
            doc.add_heading(f'Pair {i} [{pair.get("confidence", "?")}]', level=2)
            
            # Child utterance
            p1 = doc.add_paragraph()
            p1.add_run(f"Line {pair['child_line']} - Child: ").bold = True
            run1 = p1.add_run(pair['child_utterance'])
            run1.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            # Adult reformulation
            p2 = doc.add_paragraph()
            p2.add_run(f"Line {pair['adult_line']} - Adult: ").bold = True
            run2 = p2.add_run(pair['adult_utterance'])
            run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            # Correction type
            p3 = doc.add_paragraph()
            p3.add_run("Correction: ").bold = True
            p3.add_run(pair['correction_type'])
            
            # Validation info if present
            if 'validation_reason' in pair:
                p4 = doc.add_paragraph()
                p4.add_run("Validation: ").bold = True
                p4.add_run(pair['validation_reason'])
            
            doc.add_paragraph('')
    
    # Add detailed trio analysis
    if trios:
        doc.add_heading('Detailed Trio Analysis', level=1)
        for i, trio in enumerate(trios, 1):
            doc.add_heading(f'Trio {i} [{trio.get("confidence", "?")}]', level=2)
            
            # Child utterance
            p1 = doc.add_paragraph()
            p1.add_run(f"Line {trio['child_line']} - Child Error: ").bold = True
            run1 = p1.add_run(trio['child_utterance'])
            run1.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            # Adult reformulation
            p2 = doc.add_paragraph()
            p2.add_run(f"Line {trio['adult_line']} - Adult Reformulation: ").bold = True
            run2 = p2.add_run(trio['adult_reformulation'])
            run2.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            # Child self-repair
            p3 = doc.add_paragraph()
            p3.add_run(f"Line {trio['repair_line']} - Child Self-Repair: ").bold = True
            run3 = p3.add_run(trio['child_repair'])
            run3.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            
            # Correction type
            p4 = doc.add_paragraph()
            p4.add_run("Correction: ").bold = True
            p4.add_run(trio['correction_type'])
            
            # Validation info if present
            if 'validation_reason' in trio:
                p5 = doc.add_paragraph()
                p5.add_run("Validation: ").bold = True
                p5.add_run(trio['validation_reason'])
            
            doc.add_paragraph('')
    
    return doc


# Initialize session state
if "results" not in st.session_state:
    st.session_state["results"] = None
if "processing" not in st.session_state:
    st.session_state["processing"] = False
if "detection_model" not in st.session_state:
    st.session_state["detection_model"] = ""
if "validation_model" not in st.session_state:
    st.session_state["validation_model"] = ""
if "detection_prompt" not in st.session_state:
    st.session_state["detection_prompt"] = DETECTION_PROMPT
if "validation_prompt" not in st.session_state:
    st.session_state["validation_prompt"] = VALIDATION_PROMPT

# Main UI
st.title("🔍 Reformulation Detection System (Two-Pass)")
st.markdown("**Pass 1**: Small model detects ALL potential reformulations | **Pass 2**: Large model validates")

# Sidebar - Configuration
with st.sidebar:
    st.header("⚙️ Configuration")

    # BYOK Input
    user_key = st.text_input("🔑 Bring Your Own Key (Optional)", type="password", help="Enter your Ollama API key to override the default. Get one at https://ollama.com/settings/keys")
    if user_key:
        try:
            client = Client(
                host="https://ollama.com",
                headers={'Authorization': 'Bearer ' + user_key}
            )
            st.success("✓ Using custom API key")
        except Exception as e:
            st.error(f"[!] Error with custom key: {e}")

    # Model selection
    try:
        if not client:
             st.warning("⚠️ Ollama API Key not found in st.secrets. Please add it to .streamlit/secrets.toml or enter it above.")
             models = []
             detection_model = None
             validation_model = None
        else:
            models = [model.model for model in client.list().models]
            
            st.subheader("Pass 1: Detection Model")
            st.caption("Liberal detection - smaller/faster model")
            detection_model = st.selectbox(
                "Select detection model",
                models,
                index=models.index("gpt-oss:120b") if "gpt-oss:120b" in models else 0,
                key="det_model"
            )
            
            st.divider()
            
            st.subheader("Pass 2: Validation Model")
            st.caption("Strict validation - larger/smarter model")
            validation_model = st.selectbox(
                "Select validation model",
                models,
                index=models.index("gpt-oss:120b") if "gpt-oss:120b" in models else 0,
                key="val_model"
            )
        
    except Exception as e:
        st.error(f"⚠️ Error querying Ollama Cloud: {e}. Check your OLLAMA_API_KEY.")
        detection_model = None
        validation_model = None
    
    st.divider()
    
    # Chunk settings
    st.subheader("📦 Chunk Settings")
    
    chunk_size = st.slider("Chunk Size", 10, 30, 20, 
                          help="Number of lines to process at once")
    
    overlap_size = st.slider("Chunk Overlap", 0, 10, 3,
                            help="Number of lines to overlap between chunks")
    
    st.info(f"💡 {overlap_size}-line overlap prevents missing boundary pairs/trios")
    
    st.divider()
    
    # Prompt Editors
    st.subheader("📝 System Prompts")
    
    with st.expander("✏️ Pass 1: Detection Prompt", expanded=False):
        if "detection_prompt" not in st.session_state:
            st.session_state["detection_prompt"] = DETECTION_PROMPT
        
        new_detection_prompt = st.text_area(
            "Edit detection prompt (liberal approach):",
            value=st.session_state["detection_prompt"],
            height=300,
            help="This prompt guides Pass 1 detection - should be liberal",
            key="det_prompt_text"
        )
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 Update Detection", key="update_det"):
                st.session_state["detection_prompt"] = new_detection_prompt
                st.success("✅ Detection prompt updated!")
                st.rerun()
        with col2:
            if st.button("🔄 Reset Detection", key="reset_det"):
                st.session_state["detection_prompt"] = DETECTION_PROMPT
                st.success("✅ Reset to default!")
                st.rerun()
    
    with st.expander("✏️ Pass 2: Validation Prompt", expanded=False):
        if "validation_prompt" not in st.session_state:
            st.session_state["validation_prompt"] = VALIDATION_PROMPT
        
        new_validation_prompt = st.text_area(
            "Edit validation prompt (conservative approach):",
            value=st.session_state["validation_prompt"],
            height=300,
            help="This prompt guides Pass 2 validation - should be strict",
            key="val_prompt_text"
        )
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 Update Validation", key="update_val"):
                st.session_state["validation_prompt"] = new_validation_prompt
                st.success("✅ Validation prompt updated!")
                st.rerun()
        with col2:
            if st.button("🔄 Reset Validation", key="reset_val"):
                st.session_state["validation_prompt"] = VALIDATION_PROMPT
                st.success("✅ Reset to default!")
                st.rerun()
    
    st.divider()
    
    st.header("📊 Two-Pass System")
    st.markdown("""
    **Pass 1 (Detection)**:
    - Uses smaller model
    - Liberal approach
    - Catches ALL potential pairs/trios
    - Tags with confidence (HIGH/MEDIUM/LOW)
    
    **Pass 2 (Validation)**:
    - Uses larger model
    - Conservative approach
    - Filters false positives
    - Produces final validated results
    
    **Result**: Higher recall + Higher precision!
    """)
    
    st.divider()
    
    # Show current settings
    st.subheader("📋 Current Settings")
    st.markdown(f"""
    **Models:**
    - Detection: `{detection_model if detection_model else 'None'}`
    - Validation: `{validation_model if validation_model else 'None'}`
    
    **Chunk Settings:**
    - Size: {chunk_size} lines
    - Overlap: {overlap_size} lines
    
    **Prompts:**
    - Detection: {'Custom' if st.session_state.get('detection_prompt') != DETECTION_PROMPT else 'Default'}
    - Validation: {'Custom' if st.session_state.get('validation_prompt') != VALIDATION_PROMPT else 'Default'}
    """)

# Main content
tab1, tab2, tab3 = st.tabs(["📄 Upload & Process", "📊 Results", "💡 Help"])

with tab1:
    st.header("Upload Transcript")
    
    # File upload
    uploaded_file = st.file_uploader(
        "Upload transcript file",
        type=["txt", "docx"],
        help="Supported formats: .txt, .docx"
    )
    
    # Or paste text
    st.markdown("**Or paste transcript text:**")
    pasted_text = st.text_area(
        "Paste your transcript here",
        height=200,
        placeholder="1. URS: and get my pencil out first .\n2. CHI: an(d) do what ?\n..."
    )
    
    # Process button
    col1, col2 = st.columns([1, 4])
    with col1:
        process_button = st.button(
            "🚀 Process (Two-Pass)",
            disabled=not (detection_model and validation_model) or st.session_state["processing"]
        )
    with col2:
        if st.session_state["processing"]:
            st.info("Processing... Please wait")
    
    if process_button:
        if not uploaded_file and not pasted_text:
            st.error("Please upload a file or paste text")
        else:
            st.session_state["processing"] = True
            
            # Get content
            if uploaded_file:
                if uploaded_file.name.endswith('.txt'):
                    content = uploaded_file.read().decode('utf-8')
                elif uploaded_file.name.endswith('.docx') and DOCX_AVAILABLE:
                    doc = Document(uploaded_file)
                    content = '\n'.join([para.text for para in doc.paragraphs])
                else:
                    st.error("Unsupported file type")
                    st.session_state["processing"] = False
                    st.stop()
            else:
                content = pasted_text
            
            # Parse transcript
            st.subheader("📝 Parsing Transcript...")
            lines = parse_transcript(content, "txt")
            
            if len(lines) == 0:
                st.error("No utterances found! Check format: `SPEAKER: utterance`")
                st.session_state["processing"] = False
                st.stop()
            
            st.success(f"✅ Loaded {len(lines)} lines")
            
            # Create chunks
            chunks = create_chunks(lines, chunk_size, overlap_size)
            st.info(f"Processing in {len(chunks)} chunks (size: {chunk_size}, overlap: {overlap_size})")
            
            # PASS 1: DETECTION
            st.subheader("🔍 Pass 1: Detection (Liberal)")
            st.caption(f"Using {detection_model} for broad detection")
            
            all_detected_pairs = []
            all_detected_trios = []
            seen_pairs = set()
            seen_trios = set()
            chunk_data = []  # Store all chunk processing data
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            detection_container = st.container()
            
            for i, chunk in enumerate(chunks):
                progress = (i + 1) / len(chunks)
                progress_bar.progress(progress)
                status_text.text(f"Pass 1: Detecting in chunk {i+1}/{len(chunks)}")
                
                with detection_container:
                    with st.expander(f"📦 Chunk {i+1} (lines {chunk[0][0]}-{chunk[-1][0]})", expanded=False):
                        chunk_text = "\n".join([f"{num}. {text}" for num, text in chunk])
                        st.markdown("**📝 Chunk Content:**")
                        st.code(chunk_text, language="text")
                        
                        st.divider()
                        st.markdown("**🤔 Model Analysis:**")
                        
                        # Create streaming placeholder
                        stream_placeholder = st.empty()
                        status_placeholder = st.empty()
                        
                        status_placeholder.info("🧠 Model is analyzing...")
                        
                        prompt = create_analysis_prompt(chunk, "detection")
                        result = query_ollama_stream(
                            prompt,
                            detection_model,
                            st.session_state["detection_prompt"],
                            stream_placeholder
                        )
                        
                        if result['success'] and result['content']:
                            status_placeholder.success("✅ Analysis complete!")
                            
                            # Collapse the thinking/response after completion
                            with stream_placeholder.container():
                                if result['thinking']:
                                    with st.expander("💭 Model's Thinking", expanded=False):
                                        st.markdown(result['thinking'])
                                with st.expander("📄 Model's Response", expanded=False):
                                    st.code(result['content'], language="json")
                            
                            results = extract_json(result['content'])
                            chunk_pairs = results.get('pairs', [])
                            chunk_trios = results.get('trios', [])
                            
                            # Store chunk data for validation pass
                            chunk_info = {
                                'chunk_number': i + 1,
                                'lines': chunk,
                                'detected_pairs': chunk_pairs,
                                'detected_trios': chunk_trios,
                                'thinking': result['thinking'],
                                'response': result['content']
                            }
                            chunk_data.append(chunk_info)
                            
                            # Deduplicate
                            new_pairs = 0
                            new_trios = 0
                            
                            for pair in chunk_pairs:
                                pair_key = (pair['child_line'], pair['adult_line'])
                                if pair_key not in seen_pairs:
                                    seen_pairs.add(pair_key)
                                    all_detected_pairs.append(pair)
                                    new_pairs += 1
                            
                            for trio in chunk_trios:
                                trio_key = (trio['child_line'], trio['adult_line'], trio['repair_line'])
                                if trio_key not in seen_trios:
                                    seen_trios.add(trio_key)
                                    all_detected_trios.append(trio)
                                    new_trios += 1
                            
                            st.divider()
                            st.markdown("**✅ Parsed Results:**")
                            
                            if chunk_pairs or chunk_trios:
                                duplicates = (len(chunk_pairs) - new_pairs) + (len(chunk_trios) - new_trios)
                                dup_msg = f" ({duplicates} duplicates filtered)" if duplicates > 0 else ""
                                st.success(f"Detected: {new_pairs} new pairs, {new_trios} new trios{dup_msg}")
                                
                                # Show detected items
                                if new_pairs > 0:
                                    st.markdown("**🔹 New Pairs:**")
                                    for pair in chunk_pairs:
                                        if (pair['child_line'], pair['adult_line']) in seen_pairs:
                                            conf = pair.get('confidence', '?')
                                            st.markdown(f"- Lines {pair['child_line']} → {pair['adult_line']} [{conf}]: {pair['correction_type']}")
                                
                                if new_trios > 0:
                                    st.markdown("**🔸 New Trios:**")
                                    for trio in chunk_trios:
                                        if (trio['child_line'], trio['adult_line'], trio['repair_line']) in seen_trios:
                                            conf = trio.get('confidence', '?')
                                            st.markdown(f"- Lines {trio['child_line']} → {trio['adult_line']} → {trio['repair_line']} [{conf}]: {trio['correction_type']}")
                            else:
                                st.info("No potential reformulations detected")
                        else:
                            status_placeholder.error(f"❌ Error: {result.get('error', 'Unknown error')}")
            
            progress_bar.progress(100)
            status_text.text("✅ Pass 1 Complete!")
            
            st.success(f"🔍 **Pass 1 Results**: {len(all_detected_pairs)} pairs, {len(all_detected_trios)} trios detected")
            
            # Show confidence breakdown
            if all_detected_pairs or all_detected_trios:
                all_detected = all_detected_pairs + all_detected_trios
                high = sum(1 for x in all_detected if x.get('confidence') == 'HIGH')
                medium = sum(1 for x in all_detected if x.get('confidence') == 'MEDIUM')
                low = sum(1 for x in all_detected if x.get('confidence') == 'LOW')
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("HIGH Confidence", high)
                with col2:
                    st.metric("MEDIUM Confidence", medium)
                with col3:
                    st.metric("LOW Confidence", low)
            
            # PASS 2: VALIDATION
            if all_detected_pairs or all_detected_trios:
                st.divider()
                st.subheader("✅ Pass 2: Validation (Conservative)")
                st.caption(f"Using {validation_model} for strict filtering")
                
                validation_container = st.container()
                validation_chunk_data = []  # Store validation processing data
                
                # Group detected items by their original chunks for context
                items_by_chunk = {}
                for pair in all_detected_pairs:
                    for chunk_info in chunk_data:
                        for p in chunk_info['detected_pairs']:
                            if p['child_line'] == pair['child_line'] and p['adult_line'] == pair['adult_line']:
                                chunk_num = chunk_info['chunk_number']
                                if chunk_num not in items_by_chunk:
                                    items_by_chunk[chunk_num] = {
                                        'context': chunk_info['lines'],
                                        'pairs': [],
                                        'trios': []
                                    }
                                items_by_chunk[chunk_num]['pairs'].append(pair)
                
                for trio in all_detected_trios:
                    for chunk_info in chunk_data:
                        for t in chunk_info['detected_trios']:
                            if (t['child_line'] == trio['child_line'] and 
                                t['adult_line'] == trio['adult_line'] and 
                                t['repair_line'] == trio['repair_line']):
                                chunk_num = chunk_info['chunk_number']
                                if chunk_num not in items_by_chunk:
                                    items_by_chunk[chunk_num] = {
                                        'context': chunk_info['lines'],
                                        'pairs': [],
                                        'trios': []
                                    }
                                items_by_chunk[chunk_num]['trios'].append(trio)
                
                # Validate each chunk's items with context
                validated_pairs = []
                validated_trios = []
                rejected_items = []
                
                progress_bar_val = st.progress(0)
                status_text_val = st.empty()
                
                total_chunks_to_validate = len(items_by_chunk)
                
                for idx, (chunk_num, items) in enumerate(sorted(items_by_chunk.items())):
                    progress = (idx + 1) / total_chunks_to_validate
                    progress_bar_val.progress(progress)
                    status_text_val.text(f"Pass 2: Validating chunk {idx+1}/{total_chunks_to_validate}")
                    
                    with validation_container:
                        with st.expander(f"🔍 Validation - Chunk {chunk_num}", expanded=False):
                            # Show context
                            st.markdown("**📝 Conversation Context:**")
                            context_text = "\n".join([f"{num}. {text}" for num, text in items['context']])
                            st.code(context_text, language="text")
                            
                            st.divider()
                            st.markdown("**🎯 Items to Validate:**")
                            st.json({
                                'pairs': items['pairs'],
                                'trios': items['trios']
                            })
                            
                            st.divider()
                            st.markdown("**🤔 Validation Analysis:**")
                            
                            # Create streaming placeholder
                            stream_placeholder = st.empty()
                            status_placeholder = st.empty()
                            
                            status_placeholder.info("🧠 Validator is analyzing...")
                            
                            # Create validation prompt with context
                            detected_items = {
                                'pairs': items['pairs'],
                                'trios': items['trios']
                            }
                            validation_prompt_text = create_validation_prompt(detected_items, items['context'])
                            
                            result = query_ollama_stream(
                                validation_prompt_text,
                                validation_model,
                                st.session_state["validation_prompt"],
                                stream_placeholder
                            )
                            
                            if result['success'] and result['content']:
                                status_placeholder.success("✅ Validation complete!")
                                
                                # Collapse the thinking/response after completion
                                with stream_placeholder.container():
                                    if result['thinking']:
                                        with st.expander("💭 Validator's Thinking", expanded=False):
                                            st.markdown(result['thinking'])
                                    with st.expander("📄 Validator's Response", expanded=False):
                                        st.code(result['content'], language="json")
                                
                                validated_results = extract_json(result['content'])
                                
                                chunk_validated_pairs = validated_results.get('pairs', [])
                                chunk_validated_trios = validated_results.get('trios', [])
                                chunk_rejected = validated_results.get('rejected', [])
                                
                                # Store validation chunk data
                                validation_chunk_data.append({
                                    'chunk_number': chunk_num,
                                    'context': items['context'],
                                    'validated_pairs': chunk_validated_pairs,
                                    'validated_trios': chunk_validated_trios,
                                    'rejected': chunk_rejected,
                                    'thinking': result['thinking'],
                                    'response': result['content']
                                })
                                
                                validated_pairs.extend(chunk_validated_pairs)
                                validated_trios.extend(chunk_validated_trios)
                                rejected_items.extend(chunk_rejected)
                                
                                st.divider()
                                st.markdown("**✅ Validation Results:**")
                                
                                kept = len(chunk_validated_pairs) + len(chunk_validated_trios)
                                rejected = len(chunk_rejected)
                                
                                if kept > 0:
                                    st.success(f"✅ Kept: {len(chunk_validated_pairs)} pairs, {len(chunk_validated_trios)} trios")
                                if rejected > 0:
                                    st.warning(f"⚠️ Rejected: {rejected} items")
                                if kept == 0 and rejected == 0:
                                    st.info("No items to validate in this chunk")
                                    
                            else:
                                status_placeholder.error(f"❌ Validation error: {result.get('error', 'Unknown error')}")
                
                progress_bar_val.progress(100)
                status_text_val.text("✅ Pass 2 Complete!")
                
                st.success(f"✅ **Pass 2 Results**: {len(validated_pairs)} pairs, {len(validated_trios)} trios validated")
                
                if rejected_items:
                    st.warning(f"⚠️ {len(rejected_items)} items rejected as false positives")
            else:
                validated_pairs = []
                validated_trios = []
                rejected_items = []
                validation_chunk_data = []
            
            # Store results
            st.session_state["results"] = {
                "total_lines": len(lines),
                "total_chunks": len(chunks),
                "chunk_size": chunk_size,
                "overlap": overlap_size,
                "detected_pairs": len(all_detected_pairs),
                "detected_trios": len(all_detected_trios),
                "validated_pairs": len(validated_pairs),
                "validated_trios": len(validated_trios),
                "rejected_count": len(rejected_items),
                "pairs": validated_pairs,
                "trios": validated_trios,
                "rejected": rejected_items,
                "detection_model": detection_model,
                "validation_model": validation_model,
                "lines": lines,
                "detection_chunk_data": chunk_data,  # Store all detection chunk data
                "validation_chunk_data": validation_chunk_data  # Store all validation chunk data
            }
            
            st.session_state["processing"] = False
            
            # Summary
            st.success("🎉 Two-Pass Processing Complete!")
            
            col1, col2, col3, col4, col5 = st.columns(5)
            with col1:
                st.metric("Total Lines", len(lines))
            with col2:
                st.metric("Pass 1 Detected", len(all_detected_pairs) + len(all_detected_trios))
            with col3:
                st.metric("Pass 2 Validated", len(validated_pairs) + len(validated_trios))
            with col4:
                st.metric("False Positives", len(rejected_items))
            with col5:
                precision = (len(validated_pairs) + len(validated_trios)) / max(1, len(all_detected_pairs) + len(all_detected_trios)) * 100
                st.metric("Precision", f"{precision:.0f}%")

with tab2:
    st.header("Results")
    
    if st.session_state["results"] is None:
        st.info("👆 Process a transcript first to see results")
    else:
        results = st.session_state["results"]
        
        # Processing info
        st.subheader("⚙️ Processing Settings")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Detection Model", results.get("detection_model", "N/A"))
        with col2:
            st.metric("Validation Model", results.get("validation_model", "N/A"))
        with col3:
            st.metric("Chunk Size", results.get("chunk_size", "N/A"))
        with col4:
            st.metric("Overlap", results.get("overlap", "N/A"))
        
        st.divider()
        
        # Summary metrics
        st.subheader("📊 Two-Pass Summary")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Pass 1: Detected", results.get("detected_pairs", 0) + results.get("detected_trios", 0))
        with col2:
            st.metric("Pass 2: Validated", results["validated_pairs"] + results["validated_trios"])
        with col3:
            st.metric("Rejected", results.get("rejected_count", 0))
        with col4:
            total_detected = results.get("detected_pairs", 0) + results.get("detected_trios", 0)
            if total_detected > 0:
                precision = (results["validated_pairs"] + results["validated_trios"]) / total_detected * 100
                st.metric("Precision", f"{precision:.0f}%")
        
        st.divider()
        
        # Validated Pairs
        if results["pairs"]:
            st.subheader(f"✅ Validated Pairs ({len(results['pairs'])})")
            
            for i, pair in enumerate(results["pairs"], 1):
                conf = pair.get('confidence', '?')
                with st.expander(f"Pair {i} [{conf}]: Lines {pair['child_line']} → {pair['adult_line']}"):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown("**Child Utterance:**")
                        st.info(pair["child_utterance"])
                    with col2:
                        st.markdown("**Adult Reformulation:**")
                        st.success(pair["adult_utterance"])
                    
                    st.markdown(f"**Correction:** {pair['correction_type']}")
                    if 'validation_reason' in pair:
                        st.markdown(f"**Validation:** {pair['validation_reason']}")
        
        st.divider()
        
        # Validated Trios
        if results["trios"]:
            st.subheader(f"✅ Validated Trios ({len(results['trios'])})")
            
            for i, trio in enumerate(results["trios"], 1):
                conf = trio.get('confidence', '?')
                with st.expander(f"Trio {i} [{conf}]: Lines {trio['child_line']} → {trio['adult_line']} → {trio['repair_line']}"):
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.markdown("**Child Utterance:**")
                        st.info(trio["child_utterance"])
                    with col2:
                        st.markdown("**Adult Reformulation:**")
                        st.warning(trio["adult_reformulation"])
                    with col3:
                        st.markdown("**Child Self-Repair:**")
                        st.success(trio["child_repair"])
                    
                    st.markdown(f"**Correction:** {trio['correction_type']}")
                    if 'validation_reason' in trio:
                        st.markdown(f"**Validation:** {trio['validation_reason']}")
        
        # Rejected items
        if results.get("rejected"):
            st.divider()
            st.subheader(f"⚠️ Rejected Items ({len(results['rejected'])})")
            st.caption("False positives filtered out by validation pass")
            
            with st.expander("View rejected items"):
                for i, item in enumerate(results["rejected"], 1):
                    st.markdown(f"**Rejected {i}**: Lines {item['child_line']} → {item['adult_line']}")
                    st.caption(f"Reason: {item.get('rejection_reason', 'N/A')}")
                    st.divider()
        
        # Chunk Processing Data
        if results.get("detection_chunk_data") or results.get("validation_chunk_data"):
            st.divider()
            st.subheader("🔬 Detailed Chunk Processing Data")
            st.caption("View model thinking and responses for each chunk")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if results.get("detection_chunk_data"):
                    st.markdown("**📦 Pass 1: Detection Chunks**")
                    for chunk_info in results["detection_chunk_data"]:
                        with st.expander(f"Chunk {chunk_info['chunk_number']} (Detection)", expanded=False):
                            st.markdown("**Chunk Content:**")
                            chunk_text = "\n".join([f"{num}. {text}" for num, text in chunk_info['lines']])
                            st.code(chunk_text, language="text")
                            
                            st.divider()
                            
                            if chunk_info.get('thinking'):
                                st.markdown("**💭 Model's Thinking:**")
                                st.markdown(chunk_info['thinking'])
                                st.divider()
                            
                            st.markdown("**📄 Model's Response:**")
                            st.code(chunk_info['response'], language="json")
                            
                            st.divider()
                            st.markdown("**Detected:**")
                            st.write(f"- {len(chunk_info['detected_pairs'])} pairs")
                            st.write(f"- {len(chunk_info['detected_trios'])} trios")
            
            with col2:
                if results.get("validation_chunk_data"):
                    st.markdown("**🔍 Pass 2: Validation Chunks**")
                    for chunk_info in results["validation_chunk_data"]:
                        with st.expander(f"Chunk {chunk_info['chunk_number']} (Validation)", expanded=False):
                            st.markdown("**Conversation Context:**")
                            context_text = "\n".join([f"{num}. {text}" for num, text in chunk_info['context']])
                            st.code(context_text, language="text")
                            
                            st.divider()
                            
                            if chunk_info.get('thinking'):
                                st.markdown("**💭 Validator's Thinking:**")
                                st.markdown(chunk_info['thinking'])
                                st.divider()
                            
                            st.markdown("**📄 Validator's Response:**")
                            st.code(chunk_info['response'], language="json")
                            
                            st.divider()
                            st.markdown("**Results:**")
                            st.write(f"✅ Kept: {len(chunk_info['validated_pairs'])} pairs, {len(chunk_info['validated_trios'])} trios")
                            st.write(f"⚠️ Rejected: {len(chunk_info['rejected'])} items")
        
        st.divider()
        
        # Download results
        st.subheader("💾 Export Results")
        
        json_data = {k: v for k, v in results.items() if k != 'lines'}
        json_str = json.dumps(json_data, indent=2)
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button(
                label="📥 Download JSON",
                data=json_str,
                file_name="reformulation_results_validated.json",
                mime="application/json"
            )
        
        with col2:
            # Create CSV
            csv_lines = ["Type,Child Line,Child Utterance,Adult Line,Adult Utterance,Repair Line,Child Repair,Correction Type,Confidence\n"]
            
            for pair in results["pairs"]:
                conf = pair.get('confidence', '')
                csv_lines.append(f"Pair,{pair['child_line']},\"{pair['child_utterance']}\",{pair['adult_line']},\"{pair['adult_utterance']}\",,,\"{pair['correction_type']}\",{conf}\n")
            
            for trio in results["trios"]:
                conf = trio.get('confidence', '')
                csv_lines.append(f"Trio,{trio['child_line']},\"{trio['child_utterance']}\",{trio['adult_line']},\"{trio['adult_reformulation']}\",{trio['repair_line']},\"{trio['child_repair']}\",\"{trio['correction_type']}\",{conf}\n")
            
            csv_str = "".join(csv_lines)
            
            st.download_button(
                label="📥 Download CSV",
                data=csv_str,
                file_name="reformulation_results_validated.csv",
                mime="text/csv"
            )
        
        with col3:
            # Create highlighted DOCX
            if DOCX_AVAILABLE and "lines" in results:
                with st.spinner("Generating DOCX..."):
                    doc = create_highlighted_docx(
                        results["lines"],
                        results["pairs"],
                        results["trios"]
                    )
                    
                    if doc:
                        docx_buffer = BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        
                        st.download_button(
                            label="📥 Download DOCX",
                            data=docx_buffer.getvalue(),
                            file_name="reformulation_analysis_validated.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            else:
                if not DOCX_AVAILABLE:
                    st.info("Install python-docx:\n`pip install python-docx`")

with tab3:
    st.header("💡 Two-Pass System Guide")
    
    st.markdown("""
    ## How It Works
    
    ### Pass 1: Detection (Liberal)
    
    **Goal**: Cast a wide net - catch ALL potential reformulations
    
    **Model**: Small, fast model (e.g., `gpt-oss:20b`)
    
    **Approach**: Liberal detection
    - Include anything that *might* be a reformulation
    - Tag with confidence: HIGH, MEDIUM, or LOW
    - Better to over-detect than miss reformulations
    
    **Result**: High recall (catches most true reformulations)
    
    ---
    
    ### Pass 2: Validation (Conservative)
    
    **Goal**: Filter false positives - keep only TRUE reformulations
    
    **Model**: Large, smart model (e.g., `deepseek-r1`, `llama3:70b`)
    
    **Approach**: Strict validation
    - Review each detected item
    - Apply strict reformulation criteria
    - Reject expansions, interpretations, topic changes
    
    **Result**: High precision (validated results are accurate)
    
    ---
    
    ## Benefits
    
    ✅ **Higher Recall**: Small model catches edge cases
    ✅ **Higher Precision**: Large model filters false positives
    ✅ **Cost Effective**: Small model does bulk work
    ✅ **Quality Control**: Large model only validates
    ✅ **Transparency**: See what was rejected and why
    
    ---
    
    ## Recommended Models
    
    **Pass 1 (Detection)**:
    - `gpt-oss:20b` - Fast, good at pattern detection
    - `llama3:8b` - Very fast, decent quality
    - `qwen2.5:7b` - Fast alternative
    
    **Pass 2 (Validation)**:
    - `deepseek-r1:latest` - Best reasoning ability
    - `llama3:70b` - High quality validation
    - `qwen2.5:72b` - Strong alternative
    
    ---
    
    ## Settings
    
    **Chunk Size**: 20 lines (recommended)
    **Chunk Overlap**: 3 lines (prevents missing boundary pairs)
    
    ---
    
    ## Expected Results
    
    - **Pass 1**: May detect 50-100 potential items
    - **Pass 2**: Validates 20-40 true reformulations
    - **Precision**: 60-80% (depends on data)
    - **Final Rate**: 1-3% of lines (typical)
    
    ---
    
    ## Output Files
    
    1. **JSON**: All validated pairs/trios + rejection data
    2. **CSV**: Spreadsheet format with confidence tags
    3. **DOCX**: Highlighted transcript (yellow=pairs, green=trios)
    
    The DOCX file shows:
    - Full annotated transcript
    - Yellow highlights for pairs
    - Green highlights for trios
    - Confidence tags [HIGH/MEDIUM/LOW]
    - Detailed analysis for each reformulation
    """)

# Footer
st.divider()
st.markdown("Two-Pass Reformulation Detection System | Pass 1: Liberal Detection → Pass 2: Strict Validation")
